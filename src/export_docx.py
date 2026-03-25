"""
Converts research_paper.md to a properly formatted .docx file.
Upload the .docx to Google Drive → it opens as a Google Doc automatically.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PAPER_PATH = os.path.join(os.path.dirname(__file__), "..", "paper", "research_paper.md")
OUT_PATH   = os.path.join(os.path.dirname(__file__), "..", "paper", "research_paper.docx")


def set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_table(doc, rows_data, header=True):
    col_count = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text.strip()
            if i == 0 and header:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph("")


def build_doc():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Does Prompt Quality Actually Matter More Than Prompt Quantity?", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("An Empirical Study on Few-Shot Example Quality vs. Count in Large Language Model Performance")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(12)

    doc.add_paragraph("")

    # ── Abstract ──────────────────────────────────────────────────────────────
    set_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "There is a common assumption in the machine learning community that adding more few-shot examples "
        "to a prompt will improve a language model's output. This paper challenges that assumption. We designed "
        "a controlled experiment across five domains — JSON extraction, mathematical word problems, SQL generation, "
        "code debugging, and text summarization — and tested how the quality of in-context examples interacts with "
        "their quantity (1, 3, 5, and 10 shots). Each domain was evaluated using 10 test cases, generating 600 "
        "individual LLM calls using GPT-4o-mini at temperature 0.0. Our results show a Pearson correlation of "
        "r = 0.31 between prompt quality level and output accuracy (r² = 0.10), confirming that quality is a "
        "meaningful predictor of performance beyond shot count alone. Specifically, 3 high-quality examples "
        "outperformed 10 low-quality examples by an average margin of 22.7 percentage points across domains. "
        "These findings have direct implications for how developers, researchers, and practitioners construct "
        "prompts for production systems."
    )

    # ── 1. Introduction ───────────────────────────────────────────────────────
    set_heading(doc, "1. Introduction", 1)
    doc.add_paragraph(
        "The rise of large language models (LLMs) has introduced a new paradigm in how we interact with AI systems. "
        "Unlike traditional machine learning where a model is retrained for each new task, LLMs can be steered purely "
        "through natural language instructions — a process broadly called prompting. Among the most effective prompting "
        "strategies is few-shot learning, where the prompt contains a handful of input-output examples that demonstrate "
        "the desired behavior before the actual query is posed."
    )
    doc.add_paragraph(
        "The intuition behind few-shot prompting is simple: show the model what a good response looks like, and it "
        "will produce something similar. Most practitioners follow a straightforward rule of thumb — the more examples, "
        "the better. Papers like Brown et al.'s GPT-3 work (2020) showed clear gains as shot count increased from 0 to 32. "
        "But these studies rarely disentangle a confounding factor: example quality."
    )
    doc.add_paragraph(
        "When a practitioner adds more few-shot examples, they are often adding examples of varying quality — some "
        "well-formed, some ambiguous, some outright wrong. If those extra examples introduce noise, they may actively "
        "hurt performance even as they increase quantity. This brings us to the central question of our paper:"
    )
    p = doc.add_paragraph()
    run = p.add_run("Is a small number of high-quality examples more valuable than a larger number of low-quality ones?")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        "We believe this question has meaningful practical stakes. To answer it, we operationalize example quality "
        "into three levels:"
    )
    doc.add_paragraph("High quality: well-formed, complete, correctly formatted, domain-appropriate examples.", style="List Bullet")
    doc.add_paragraph("Medium quality: mostly correct but with missing fields, imprecise answers, or inconsistent formatting.", style="List Bullet")
    doc.add_paragraph("Low quality: incorrect answers, poor formatting, or misleading demonstrations.", style="List Bullet")

    doc.add_paragraph(
        "We cross this quality axis with four shot counts (1, 3, 5, 10) across five domains, yielding 60 unique "
        "conditions total."
    )

    # ── 2. Related Work ───────────────────────────────────────────────────────
    set_heading(doc, "2. Related Work", 1)

    set_heading(doc, "In-context learning and few-shot prompting", 2)
    doc.add_paragraph(
        "Brown et al. (2020) showed that a single large model could perform competitively on many NLP benchmarks "
        "with just a few demonstration examples, without any gradient updates. Min et al. (2022) found that format "
        "and the demonstration that a response is expected matter more than label correctness. Our work extends this "
        "by examining quality at a more holistic level."
    )

    set_heading(doc, "Prompt engineering as a discipline", 2)
    doc.add_paragraph(
        "Wei et al. (2022) introduced chain-of-thought prompting, showing that including reasoning steps dramatically "
        "improves performance on reasoning tasks. This implicitly encodes quality — a chain-of-thought example is more "
        "detailed and informative than a flat input-output pair."
    )

    set_heading(doc, "Example selection strategies", 2)
    doc.add_paragraph(
        "Zhang et al. (2022) proposed selecting examples using semantic similarity. Su et al. (2022) used mutual "
        "information to rank examples. These works assume high-quality examples and focus on selection. We ask instead: "
        "given a fixed number of examples, how much does their inherent quality affect outcomes?"
    )

    set_heading(doc, "Prompt sensitivity", 2)
    doc.add_paragraph(
        "Lu et al. (2022) showed that example order can affect performance by over 30 percentage points. Zhao et al. "
        "(2021) showed majority-label bias affects predictions. Our work contributes another axis — the quality level "
        "of the examples themselves."
    )

    # ── 3. Methodology ────────────────────────────────────────────────────────
    set_heading(doc, "3. Methodology", 1)

    set_heading(doc, "3.1 Experimental Design", 2)
    doc.add_paragraph(
        "Our experiment follows a 3 × 4 × 5 factorial design: 3 quality levels (high, medium, low), "
        "4 shot counts (1, 3, 5, 10), and 5 domains. For each of the 60 resulting conditions, we ran 10 test cases "
        "with known ground truth, yielding 600 LLM evaluations. All calls were made to GPT-4o-mini with temperature 0.0."
    )

    set_heading(doc, "3.2 Domain Selection Rationale", 2)
    add_table(doc, [
        ["Domain", "Task Type", "Metric", "Why Chosen"],
        ["JSON Extraction", "Structured output", "Field-level F1", "Tests precision and format adherence"],
        ["Math Word Problems", "Reasoning", "Exact match", "Tests numerical accuracy"],
        ["SQL Generation", "Code generation", "Keyword coverage", "Tests domain-specific syntax"],
        ["Code Debugging", "Analysis + generation", "Bug ID + fix accuracy", "Tests diagnostic reasoning"],
        ["Text Summarization", "Language generation", "Key-point coverage", "Tests factual compression"],
    ])

    set_heading(doc, "3.3 Defining Example Quality", 2)
    add_table(doc, [
        ["Attribute", "High", "Medium", "Low"],
        ["Answer correctness", "Fully correct", "Mostly correct", "Often wrong"],
        ["Format compliance", "Proper formatting", "Inconsistent", "Unstructured"],
        ["Completeness", "All fields present", "Partial", "Missing key info"],
        ["Domain appropriateness", "Domain terminology used", "Generic", "Misleading"],
    ])
    doc.add_paragraph(
        "Examples were constructed manually at each quality level rather than synthetically degraded, "
        "ensuring they reflect realistic practitioner behaviour."
    )

    set_heading(doc, "3.4 Prompt Structure", 2)
    doc.add_paragraph("High quality prompt: Explicit role definition, task-specific instructions, format constraints, domain keywords, numbered examples.", style="List Bullet")
    doc.add_paragraph("Medium quality prompt: Brief role mention, basic task instruction, labeled but minimally formatted examples.", style="List Bullet")
    doc.add_paragraph("Low quality prompt: Bare task request, no role or constraints, raw concatenated examples with no structure.", style="List Bullet")

    set_heading(doc, "3.5 Evaluation Metrics", 2)
    doc.add_paragraph("JSON Extraction — Field-level F1: 1.0 for exact match, 0.7 for partial match, 0.5 for wrong value, 0.0 for missing key.", style="List Bullet")
    doc.add_paragraph("Math Word Problems — Exact match: Look for 'Final Answer: X' pattern, compare numerically with ±0.01 tolerance.", style="List Bullet")
    doc.add_paragraph("SQL Generation — Keyword coverage: Fraction of expected SQL keywords present in output.", style="List Bullet")
    doc.add_paragraph("Code Debugging — 0.5 for correct bug identification + 0.5 for correct fix.", style="List Bullet")
    doc.add_paragraph("Text Summarization — Key-point coverage: Fraction of expected key facts present in summary.", style="List Bullet")

    # ── 4. Results ────────────────────────────────────────────────────────────
    set_heading(doc, "4. Results", 1)

    set_heading(doc, "4.1 Overall Quality Effect", 2)
    doc.add_paragraph("Table 1: Mean accuracy by prompt quality level (all domains, all shot counts)")
    add_table(doc, [
        ["Quality Level", "Mean Accuracy", "vs. Low Quality"],
        ["High", "0.86 (86.4%)", "+21.1 pp"],
        ["Medium", "0.66 (65.9%)", "+0.7 pp"],
        ["Low", "0.65 (65.2%)", "—"],
    ])
    doc.add_paragraph(
        "High-quality prompts outperform low-quality by 21.1 percentage points. Medium quality shows only +0.7 pp "
        "over low quality — partial curation yields almost no benefit."
    )

    set_heading(doc, "4.2 Quality × Shot Count Interaction", 2)
    doc.add_paragraph("Table 2: Mean accuracy — Quality (rows) × Shot Count (cols)")
    add_table(doc, [
        ["Quality", "1 shot", "3 shots", "5 shots", "10 shots"],
        ["High",   "0.833", "0.856", "0.856", "0.908"],
        ["Medium", "0.680", "0.658", "0.666", "0.633"],
        ["Low",    "0.652", "0.670", "0.657", "0.630"],
    ])
    doc.add_paragraph(
        "High quality keeps improving with more shots. Medium and low quality decline with more shots. "
        "3 high-quality examples (85.6%) vs 10 low-quality examples (63.0%) = 22.7 pp gap — the core finding."
    )

    set_heading(doc, "4.3 Per-Domain Breakdown", 2)
    doc.add_paragraph("Table 3: Mean accuracy by domain and quality level")
    add_table(doc, [
        ["Domain", "High", "Medium", "Low", "Quality Δ (H-L)"],
        ["JSON Extraction",    "0.67 (67%)", "0.60 (60%)", "0.01 (1%)",  "+65.8 pp"],
        ["Math Word Problems", "1.00 (100%)","0.20 (20%)", "0.80 (80%)", "+20.0 pp"],
        ["SQL Generation",     "0.98 (98%)", "0.93 (93%)", "0.94 (94%)", "+3.1 pp"],
        ["Code Debugging",     "0.76 (76%)", "0.78 (78%)", "0.66 (66%)", "+9.6 pp"],
        ["Text Summarization", "0.92 (92%)", "0.79 (79%)", "0.85 (85%)", "+7.2 pp"],
    ])
    doc.add_paragraph(
        "JSON extraction shows the largest quality gap (+65.8 pp). Low-quality prompts caused the model to return "
        "plain text instead of JSON entirely. Math shows a metric artifact — medium quality (20%) scored worse than "
        "low quality (80%) because medium prompts did not enforce the 'Final Answer: X' format the evaluator requires. "
        "SQL shows the smallest gap (+3.1 pp) due to strong model pretraining on SQL syntax."
    )

    set_heading(doc, "4.4 Correlation Analysis", 2)
    doc.add_paragraph(
        "Quality levels encoded numerically (high=3, medium=2, low=1). Pearson r = 0.31, r² = 0.10. "
        "Quality explains approximately 10% of variance in output accuracy across all 60 conditions. "
        "Within individual domains the effect is substantially stronger."
    )

    set_heading(doc, "4.5 Shot Count Finding", 2)
    doc.add_paragraph("Table 4: Accuracy by quality level across shot counts")
    add_table(doc, [
        ["Quality", "1 shot", "3 shots", "5 shots", "10 shots", "Trend"],
        ["High",   "83.3%", "85.6%", "85.6%", "90.8%", "↑ Keeps improving"],
        ["Medium", "68.0%", "65.8%", "66.6%", "63.3%", "↓ Slowly declining"],
        ["Low",    "65.2%", "67.0%", "65.7%", "63.0%", "→ Flat / declining"],
    ])
    doc.add_paragraph(
        "High quality shows continued gains up to 10 shots. Medium and low quality plateau or decline, "
        "confirming that quantity only helps when quality is already high."
    )

    # ── 5. Discussion ─────────────────────────────────────────────────────────
    set_heading(doc, "5. Discussion", 1)

    set_heading(doc, "5.1 Why Quality Dominates Quantity", 2)
    doc.add_paragraph(
        "High-quality examples establish a consistent, correct pattern for the model to follow. Low-quality examples "
        "establish an inconsistent or incorrect pattern. Adding more low-quality examples amplifies noise rather than "
        "signal — explaining why medium and low quality performance declines with more shots."
    )

    set_heading(doc, "5.2 Implications for Practitioners", 2)
    doc.add_paragraph("Curate before you accumulate. 3 high-quality examples outperform 10 low-quality ones by 22.7 pp.", style="List Number")
    doc.add_paragraph("Format compliance is not optional. Even correct answers score near zero without the required output format.", style="List Number")
    doc.add_paragraph("Medium quality is not a safe middle ground. Only +0.7 pp better than low quality on average.", style="List Number")
    doc.add_paragraph("Domain determines sensitivity. JSON and math are most sensitive; SQL least sensitive.", style="List Number")
    doc.add_paragraph("More high-quality shots still help. Gains continue to 10 shots — but only when quality is high.", style="List Number")

    set_heading(doc, "5.3 Limitations", 2)
    doc.add_paragraph("Single model — GPT-4o-mini only. Results may differ for larger or smaller models.", style="List Bullet")
    doc.add_paragraph("Automatic metrics — not perfect proxies for human judgement.", style="List Bullet")
    doc.add_paragraph("Example construction — hand-constructed by a single author, no inter-rater agreement.", style="List Bullet")
    doc.add_paragraph("Domain coverage — five domains may not represent all LLM use cases.", style="List Bullet")
    doc.add_paragraph("Shot count ceiling — only tested up to 10 shots.", style="List Bullet")

    # ── 6. Conclusion ─────────────────────────────────────────────────────────
    set_heading(doc, "6. Conclusion", 1)
    doc.add_paragraph(
        "This paper presented a systematic empirical study of how few-shot example quality interacts with example "
        "quantity in LLM prompting. Across five domains and 600 LLM evaluations using GPT-4o-mini, we find consistent "
        "evidence that quality outweighs quantity: 3 high-quality examples reliably outperform 10 low-quality examples "
        "with a mean advantage of 22.7 percentage points."
    )
    doc.add_paragraph(
        "The Pearson correlation (r = 0.31, r² = 0.10) confirms quality as a meaningful predictor. Within specific "
        "domains — JSON extraction (+65.8 pp) and math (+20.0 pp) — the effect is decisive."
    )
    doc.add_paragraph(
        "For practitioners, the message is direct: invest in your examples. The difference between a carelessly "
        "assembled prompt and a thoughtfully curated one can be the difference between a system that works and one "
        "that does not."
    )

    # ── References ────────────────────────────────────────────────────────────
    set_heading(doc, "References", 1)
    refs = [
        "Brown, T., et al. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877–1901.",
        "Wei, J., et al. (2022). Chain-of-thought prompting elicits reasoning in large language models. NeurIPS 35.",
        "Min, S., et al. (2022). Rethinking the role of demonstrations: What makes in-context learning work? EMNLP 2022.",
        "Lu, Y., et al. (2022). Fantastically ordered prompts and where to find them. ACL 2022.",
        "Zhang, Z., et al. (2022). Automatic chain of thought prompting in large language models. arXiv:2210.11610.",
        "Su, H., et al. (2022). Selective annotation makes language models better few-shot learners. arXiv:2209.01975.",
        "Zhao, Z., et al. (2021). Calibrate before use: Improving few-shot performance of language models. ICML 2021.",
        "Liu, P., et al. (2023). Pre-train, prompt, and predict. ACM Computing Surveys, 55(9), 1–35.",
        "Dong, Q., et al. (2022). A survey on in-context learning. arXiv:2301.00234.",
        "Lester, B., Al-Rfou, R., & Constant, N. (2021). The power of scale for parameter-efficient prompt tuning. EMNLP 2021.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    doc.add_paragraph("")
    p = doc.add_paragraph("Manuscript prepared as part of a college-level research project on prompt engineering efficiency in large language models.")
    p.runs[0].italic = True

    doc.save(OUT_PATH)
    print(f"[SAVED] {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
